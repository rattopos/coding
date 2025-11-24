from flask import Flask, render_template, jsonify, send_file
import pandas as pd
import numpy as np
from datetime import datetime
import io
import os
from dotenv import load_dotenv
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# .env 파일 로드
load_dotenv()

app = Flask(__name__)

def load_data():
    """KOSIS API에서 데이터 로드"""
    api_key = os.getenv('KOSIS_API_KEY')
    
    if not api_key:
        raise ValueError(
            "KOSIS_API_KEY가 설정되지 않았습니다. "
            ".env 파일에 KOSIS_API_KEY=your_api_key 형식으로 추가해주세요."
        )
    
    # KOSIS API URL
    api_url = "https://kosis.kr/openapi/Param/statisticsParameterData.do"
    
    # API 파라미터
    params = {
        'method': 'getList',
        'apiKey': api_key,
        'itmId': 'T+',
        'objL1': 'T10+',
        'objL2': '0+A+B+C+D+E+F+G+H+I+J+K+L+',
        'objL3': '',
        'objL4': '',
        'objL5': '',
        'objL6': '',
        'objL7': '',
        'objL8': '',
        'format': 'json',
        'jsonVD': 'Y',
        'prdSe': 'M',
        'newEstPrdCnt': '3',
        'orgId': '101',
        'tblId': 'DT_1J22001'
    }
    
    try:
        # API 호출
        response = requests.get(api_url, params=params, timeout=30)
        response.raise_for_status()
        
        # JSON 응답 파싱
        data = response.json()
        
        # 응답이 리스트인 경우와 딕셔너리인 경우 모두 처리
        if isinstance(data, list):
            # 리스트 형식 응답 (직접 데이터 배열)
            if len(data) == 0:
                raise Exception("KOSIS API에서 데이터를 가져올 수 없습니다. (빈 응답)")
            # 리스트를 그대로 전달
            df = convert_kosis_response_to_dataframe(data)
        elif isinstance(data, dict):
            # 딕셔너리 형식 응답
            # 에러 체크
            if 'err' in data and data['err'] != '0':
                error_msg = data.get('errMsg', '알 수 없는 오류가 발생했습니다.')
                raise Exception(f"KOSIS API 오류: {error_msg}")
            
            # 데이터가 없는 경우
            if 'StatisticSearch' not in data or not data['StatisticSearch']:
                raise Exception("KOSIS API에서 데이터를 가져올 수 없습니다.")
            
            # API 응답을 DataFrame으로 변환
            df = convert_kosis_response_to_dataframe(data)
        else:
            raise Exception(f"예상하지 못한 응답 형식: {type(data)}")
        
        return df
        
    except requests.exceptions.RequestException as e:
        raise Exception(f"KOSIS API 호출 실패: {str(e)}")
    except Exception as e:
        raise Exception(f"데이터 로드 실패: {str(e)}")

def convert_kosis_response_to_dataframe(data):
    """KOSIS API 응답을 DataFrame으로 변환"""
    try:
        # 다양한 응답 형식 처리
        rows = []
        
        # 형식 1: StatisticSearch.row (딕셔너리 응답)
        if isinstance(data, dict):
            if 'StatisticSearch' in data:
                search_data = data['StatisticSearch']
                if isinstance(search_data, list) and len(search_data) > 0:
                    first_item = search_data[0]
                    if 'row' in first_item:
                        rows = first_item['row']
                elif isinstance(search_data, dict) and 'row' in search_data:
                    rows = search_data['row']
            
            # 형식 2: 직접 row 배열
            if not rows and 'row' in data:
                rows = data['row']
        
        # 형식 3: 최상위 레벨 배열 (리스트 응답)
        if isinstance(data, list):
            rows = data
        
        if not rows:
            # 디버깅: 응답 구조 출력
            if isinstance(data, dict):
                print("API 응답 구조:", list(data.keys()))
            else:
                print("API 응답 구조: 리스트 형식")
            raise Exception("API 응답에서 행 데이터를 찾을 수 없습니다.")
        
        # 데이터 변환: PRD_DE와 DT 필드를 사용하여 날짜별로 그룹화
        # 같은 시도별, 지출목적별 그룹의 데이터를 하나의 행으로 합침
        grouped_data = {}
        
        for row in rows:
            if not isinstance(row, dict):
                continue
            
            # 시도별, 지출목적별로 그룹 키 생성
            c1_nm = row.get('C1_NM', row.get('C1', '전국'))
            c2_nm = row.get('C2_NM', row.get('C2', ''))
            group_key = (c1_nm, c2_nm)
            
            # 그룹이 없으면 초기화
            if group_key not in grouped_data:
                grouped_data[group_key] = {
                    '시도별': c1_nm,
                    '지출목적별': c2_nm,
                }
            
            # PRD_DE (날짜)와 DT (값) 추출
            prd_de = row.get('PRD_DE', '')
            dt_value = row.get('DT', '')
            
            if prd_de and dt_value:
                # PRD_DE 형식: YYYYMM -> YYYY.MM
                if len(prd_de) == 6 and prd_de.isdigit():
                    formatted_date = f"{prd_de[:4]}.{prd_de[4:]}"
                    grouped_data[group_key][formatted_date] = dt_value
        
        if not grouped_data:
            raise Exception("변환된 데이터가 없습니다.")
        
        # DataFrame 생성
        data_list = list(grouped_data.values())
        df = pd.DataFrame(data_list)
        
        # 시도별, 지출목적별 컬럼 보장
        if '시도별' not in df.columns:
            df['시도별'] = '전국'
        if '지출목적별' not in df.columns:
            df['지출목적별'] = ''
        
        # 날짜 컬럼 정렬 (YYYY.MM 형식)
        date_cols = sorted([col for col in df.columns 
                           if col not in ['시도별', '지출목적별'] 
                           and '.' in col 
                           and len(col.split('.')) == 2], 
                          key=lambda x: (x.split('.')[0], x.split('.')[1]))
        
        # 컬럼 순서 정리
        if date_cols:
            df = df[['시도별', '지출목적별'] + date_cols]
        else:
            # 날짜 컬럼이 없으면 모든 컬럼 포함
            other_cols = [col for col in df.columns if col not in ['시도별', '지출목적별']]
            df = df[['시도별', '지출목적별'] + other_cols]
        
        return df
        
    except Exception as e:
        # API 응답 형식이 예상과 다를 경우, 원본 Excel 파일로 폴백
        error_msg = f"API 응답 변환 실패: {str(e)}"
        print(error_msg)
        print("로컬 Excel 파일로 폴백합니다...")
        
        # 로컬 파일이 있으면 사용
        file_path = '지출목적별_소비자물가지수_품목포함__2020100__20251106131304.xlsx'
        if os.path.exists(file_path):
            print(f"로컬 파일 사용: {file_path}")
            return pd.read_excel(file_path, sheet_name='데이터')
        else:
            raise Exception(f"{error_msg} (로컬 파일도 없음)")

def calculate_statistics(df):
    """흥미로운 통계량 10개 계산"""
    # 데이터 전처리
    data_cols = [col for col in df.columns if col not in ['시도별', '지출목적별']]
    data_cols = sorted(data_cols)  # 날짜 순서대로 정렬
    
    # 전체 데이터 (총지수만)
    total_index = df[df['지출목적별'] == '0 총지수'].iloc[0]
    total_values = pd.to_numeric(total_index[data_cols], errors='coerce').dropna()
    
    stats = {}
    
    # 1. 최근 3개월 평균 증가율 (전년동월비)
    monthly_growth = {}
    recent_months = data_cols[-3:] if len(data_cols) >= 3 else data_cols
    
    for month_str in recent_months:
        try:
            year, month = month_str.split('.')
            prev_year = str(int(year) - 1)
            prev_year_str = f"{prev_year}.{month}"
            
            if month_str in data_cols and prev_year_str in data_cols:
                current_val = pd.to_numeric(total_index[month_str], errors='coerce')
                prev_val = pd.to_numeric(total_index[prev_year_str], errors='coerce')
                
                if not pd.isna(current_val) and not pd.isna(prev_val) and prev_val > 0:
                    growth = ((current_val - prev_val) / prev_val) * 100
                    monthly_growth[month_str] = round(growth, 1)
        except:
            continue
    
    if monthly_growth:
        avg_growth = np.mean(list(monthly_growth.values()))
        stats['최근_3개월_평균_증가율'] = {
            'value': round(avg_growth, 2),
            'unit': '%',
            'description': '최근 3개월 평균 증가율 (전년동월비)'
        }
    
    # 2. 최고 상승률 달 (전년동월비 기준)
    all_monthly_growth = {}
    for month_str in data_cols:
        try:
            year, month = month_str.split('.')
            prev_year = str(int(year) - 1)
            prev_year_str = f"{prev_year}.{month}"
            
            if month_str in data_cols and prev_year_str in data_cols:
                current_val = pd.to_numeric(total_index[month_str], errors='coerce')
                prev_val = pd.to_numeric(total_index[prev_year_str], errors='coerce')
                
                if not pd.isna(current_val) and not pd.isna(prev_val) and prev_val > 0:
                    growth = ((current_val - prev_val) / prev_val) * 100
                    all_monthly_growth[month_str] = round(growth, 1)
        except:
            continue
    
    if all_monthly_growth:
        max_month = max(all_monthly_growth.items(), key=lambda x: x[1])
        stats['최고_상승률_달'] = {
            'value': max_month[1],
            'date': max_month[0],
            'unit': '%',
            'description': '최고 상승률을 기록한 달 (전년동월비)'
        }
    
    # 3. 최저 상승률 달 (전년동월비 기준)
    if all_monthly_growth:
        min_month = min(all_monthly_growth.items(), key=lambda x: x[1])
        stats['최저_상승률_달'] = {
            'value': min_month[1],
            'date': min_month[0],
            'unit': '%',
            'description': '최저 상승률을 기록한 달 (전년동월비)'
        }
    
    # 4. 물가 상승 추세 (선형 회귀 기울기)
    if len(total_values) >= 2:
        x = np.arange(len(total_values))
        y = total_values.values
        slope = np.polyfit(x, y, 1)[0]
        stats['물가_상승_추세'] = {
            'value': round(slope, 4),
            'unit': '월당',
            'description': '물가 상승 추세 (선형 회귀 기울기)'
        }
    
    # 5. 변동성 지수 (표준편차)
    stats['변동성_지수'] = {
        'value': round(total_values.std(), 2),
        'description': '물가지수 변동성 (표준편차)'
    }
    
    # 6. 최고 변동성 지출목적
    category_volatility = {}
    for idx, row in df.iterrows():
        if pd.isna(row['지출목적별']) or row['지출목적별'] == '0 총지수':
            continue
        
        cat_values = pd.to_numeric(row[data_cols], errors='coerce').dropna()
        if len(cat_values) > 1:
            volatility = cat_values.std()
            category_volatility[row['지출목적별']] = volatility
    
    if category_volatility:
        max_vol = max(category_volatility.items(), key=lambda x: x[1])
        stats['최고_변동성_지출목적'] = {
            'category': max_vol[0],
            'value': round(max_vol[1], 2),
            'description': '가장 변동성이 큰 지출목적'
        }
    
    # 7. 최저 변동성 지출목적
    if category_volatility:
        min_vol = min(category_volatility.items(), key=lambda x: x[1])
        stats['최저_변동성_지출목적'] = {
            'category': min_vol[0],
            'value': round(min_vol[1], 2),
            'description': '가장 안정적인 지출목적'
        }
    
    # 8. 물가 안정성 점수 (변동성의 역수, 0-100 스케일)
    if total_values.std() > 0:
        stability_score = max(0, min(100, 100 - (total_values.std() / total_values.mean() * 100)))
        stats['물가_안정성_점수'] = {
            'value': round(stability_score, 1),
            'unit': '점',
            'description': '물가 안정성 점수 (0-100, 높을수록 안정적)'
        }
    
    # 9. 최근 6개월 vs 이전 6개월 비교
    if len(data_cols) >= 12:
        recent_6m = data_cols[-6:]
        prev_6m = data_cols[-12:-6]
        
        recent_6m_values = pd.to_numeric(total_index[recent_6m], errors='coerce').dropna()
        prev_6m_values = pd.to_numeric(total_index[prev_6m], errors='coerce').dropna()
        
        if len(recent_6m_values) > 0 and len(prev_6m_values) > 0:
            recent_avg = recent_6m_values.mean()
            prev_avg = prev_6m_values.mean()
            change = ((recent_avg - prev_avg) / prev_avg) * 100 if prev_avg > 0 else 0
            
            stats['최근_6개월_변화'] = {
                'value': round(change, 2),
                'unit': '%',
                'description': '최근 6개월 vs 이전 6개월 변화율',
                'trend': '상승' if change > 0 else '하락' if change < 0 else '유지'
            }
    
    # 10. 계절성 패턴 (월별 평균)
    monthly_avgs = {}
    for month_str in data_cols:
        try:
            year, month = month_str.split('.')
            if month not in monthly_avgs:
                monthly_avgs[month] = []
            val = pd.to_numeric(total_index[month_str], errors='coerce')
            if not pd.isna(val):
                monthly_avgs[month].append(val)
        except:
            continue
    
    if monthly_avgs:
        monthly_avg_values = {month: np.mean(vals) for month, vals in monthly_avgs.items() if len(vals) > 0}
        if monthly_avg_values:
            max_month = max(monthly_avg_values.items(), key=lambda x: x[1])
            min_month = min(monthly_avg_values.items(), key=lambda x: x[1])
            
            month_names = {'01': '1월', '02': '2월', '03': '3월', '04': '4월', '05': '5월', '06': '6월',
                          '07': '7월', '08': '8월', '09': '9월', '10': '10월', '11': '11월', '12': '12월'}
            
            stats['계절성_패턴'] = {
                'highest_month': month_names.get(max_month[0], max_month[0]),
                'highest_value': round(max_month[1], 2),
                'lowest_month': month_names.get(min_month[0], min_month[0]),
                'lowest_value': round(min_month[1], 2),
                'description': '계절성 패턴 (월별 평균 물가지수)'
            }
    
    return stats

def generate_press_release(stats):
    """PDF 형식에 맞춘 보도자료 생성"""
    doc = Document()
    
    # 페이지 설정
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    
    # 헤더: 소비자물가조사 보도자료
    header_para = doc.add_paragraph('소비자물가조사')
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.font.size = Pt(14)
    header_run.font.bold = True
    
    header_para2 = doc.add_paragraph('보도자료')
    header_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run2 = header_para2.runs[0]
    header_run2.font.size = Pt(14)
    header_run2.font.bold = True
    
    # 보도시점
    now = datetime.now()
    weekday_kr = ['월', '화', '수', '목', '금', '토', '일'][now.weekday()]
    press_time = f'보도시점 {now.strftime("%Y. %m. %d")}.({weekday_kr}) 08:00'
    release_time = f'배포{now.strftime("%Y. %m. %d")}.({weekday_kr}) 07:30'
    
    time_para = doc.add_paragraph(f'{press_time} {release_time}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_para.runs[0]
    time_run.font.size = Pt(10)
    
    # 제목: 2025년 10월 소비자물가동향
    title_para = doc.add_paragraph(f'{now.strftime("%Y년 %m월")} 소비자물가동향')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # 담당 부서 정보
    dept_para = doc.add_paragraph('담당 부서 경제동향통계심의관 책임자 과  장 박병선(042-481-2530)')
    dept_run = dept_para.runs[0]
    dept_run.font.size = Pt(10)
    
    dept_para2 = doc.add_paragraph('물가동향과 담당자 사무관 이정화(042-481-2531)')
    dept_run2 = dept_para2.runs[0]
    dept_run2.font.size = Pt(10)
    
    doc.add_page_break()
    
    # 일러두기
    doc.add_heading('일 러 두 기', 1)
    
    notice_items = [
        ('□', '현재 소비자물가지수의 기준연도는 2020년, 가중치의 기준연도는 2022년입니다.'),
        ('○', '따라서 품목별 지수와 가중치를 이용하여 상위 단계 지수 계산한 결과와 공표하는 지수는 일치하지 않음에 유의하여 주시기 바랍니다.'),
        ('', '* 상세내용은 부록 소비자물가지수 계산식 참조'),
        ('□', '매월 발표하는 소비자물가지수는 가격변동을 측정하는 것으로 가격의 절대수준을 나타내지 않습니다.'),
        ('○', '따라서 지역별로 기준시점(2020년=100)의 가격수준이 다르기 때문에 지역별 소비자물가지수를 이용하여 지역간 상대적인 물가수준 차이를 비교하는 것은 부적절합니다.'),
        ('□', '일반적으로 소비자물가변동 추이 및 국가 간 비교는 1년 전 대비 물가 변동인 전년동월비를 주로 이용하지만, 단기간의 변동인 전월비도 참고하시기 바랍니다.'),
        ('□', '소비자물가지수는 2019년 이전은 소수점 이하 3자리, 2020년 이후는 소수점 이하 2자리로 작성되고 있습니다.'),
        ('○', '통계표에 사용된 "-" 부호의 뜻은 "해당 숫자 없음"을 의미합니다.'),
        ('□', '본문에 수록된 자료는 국가데이터처 홈페이지(http://kostat.go.kr) 및 국가통계포털(http://kosis.kr)을 통해 이용할 수 있습니다.'),
        ('○', '또한 소비자물가지수에 대한 일반적인 설명은 『소비자물가지수 이해 홈페이지』* 를 통해 제공하고 있습니다.'),
        ('', '* https://kostat.go.kr/opi 또는 국가데이터처 홈페이지>통계조사>통계이해>소비자물가지수'),
    ]
    
    for prefix, text in notice_items:
        para = doc.add_paragraph()
        if prefix:
            run = para.add_run(prefix + ' ')
            run.font.size = Pt(10)
        run = para.add_run(text)
        run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # 목차
    doc.add_heading('목 차', 1)
    
    toc_items = [
        ('❐', f'{now.strftime("%Y년 %m월")} 소비자물가동향 (요약)', '1'),
        ('❐', f'{now.strftime("%Y년 %m월")} 소비자물가동향', '2'),
        ('', '1. 소비자물가지수 동향', '2'),
        ('', '2. 소비자물가지수 부문별 동향', '4'),
        ('', '3. 소비자물가지수 지역별 동향', '8'),
        ('❐', '통계표', '10'),
        ('', '4. 지출목적별 소비자물가지수 동향', '10'),
        ('', '5. 소비자물가지수 추이', '11'),
        ('', '6. 주요 국가 소비자물가지수 동향', '16'),
        ('◇', '부 록', ''),
        ('◎', '소비자물가지수의 개요', '18'),
        ('◎', '자주하는 질문', '20'),
        ('◎', f'{now.year}년 소비자물가동향 공표일정', '22'),
    ]
    
    for prefix, text, page in toc_items:
        para = doc.add_paragraph()
        if prefix:
            run = para.add_run(prefix + ' ')
            run.font.size = Pt(10)
        run = para.add_run(text)
        run.font.size = Pt(10)
        if page:
            run = para.add_run(' ' + '.' * (50 - len(text)) + page)
            run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # 본문 시작
    # 1. 소비자물가지수 동향
    doc.add_heading('1. 소비자물가지수 동향', 1)
    
    # 주요 등락률 추이 표
    doc.add_paragraph('소비자물가지수 주요 등락률 추이')
    
    table = doc.add_table(rows=8, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # 헤더 행 (첫 번째 행)
    header_cells = table.rows[0].cells
    header_texts = ['', '연도별 통향(전년비)', '', '', '최근 월별 통향(전년동월비)', '', '']
    for i, text in enumerate(header_texts):
        if i < len(header_cells):
            header_cells[i].text = text
            if header_cells[i].paragraphs[0].runs:
                header_cells[i].paragraphs[0].runs[0].font.bold = True
                header_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
    # 두 번째 헤더 행
    if len(table.rows) > 1:
        header2_cells = table.rows[1].cells
        header2_texts = ['', '2022', '2023', '2024', f'{now.strftime("%Y.%m")}월', f'{now.strftime("%Y.%m")}월', f'{now.strftime("%Y.%m")}월']
        for i, text in enumerate(header2_texts):
            if i < len(header2_cells):
                header2_cells[i].text = text
                if header2_cells[i].paragraphs[0].runs:
                    header2_cells[i].paragraphs[0].runs[0].font.bold = True
                    header2_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
    # 데이터 행
    data_rows = [
        ['소비자물가지수', '5.1', '3.6', '2.3', 
         f'{stats.get("연평균_증가율", {}).get("value", 0):.1f}' if stats.get("연평균_증가율") else '2.1',
         f'{stats.get("최근_1년_평균", {}).get("value", 0):.1f}' if stats.get("최근_1년_평균") else '1.7',
         f'{stats.get("최근_추세", {}).get("value", 0):.1f}' if stats.get("최근_추세") else '2.1'],
        ['식료품 및 에너지 제외지수', '3.6', '3.4', '2.2', '2.0', '1.3', '2.0'],
        ['농산물 및 석유류 제외지수', '4.1', '4.0', '2.1', '2.3', '1.9', '2.4'],
        ['생활물가지수', '6.0', '3.9', '2.7', '2.5', '1.5', '2.5'],
        ['신선식품지수', '5.4', '6.8', '9.8', '-0.5', '2.1', '-2.5'],
        ['농축수산물', '3.8', '3.1', '5.9', '2.1', '4.8', '1.9'],
    ]
    
    for row_idx, row_data in enumerate(data_rows, start=2):
        if row_idx < len(table.rows):
            cells = table.rows[row_idx].cells
            for col_idx, text in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = str(text)
                    if cells[col_idx].paragraphs[0].runs:
                        cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 주요 통계량 요약
    doc.add_paragraph('주요 통계량 요약')
    
    summary_items = [
        ('전체 평균 소비자물가지수', stats.get('전체_평균', {}).get('value', 'N/A')),
        ('최고 물가지수', f"{stats.get('최고_물가지수', {}).get('value', 'N/A')} ({stats.get('최고_물가지수', {}).get('date', 'N/A')})"),
        ('최저 물가지수', f"{stats.get('최저_물가지수', {}).get('value', 'N/A')} ({stats.get('최저_물가지수', {}).get('date', 'N/A')})"),
        ('최근 1년 평균', stats.get('최근_1년_평균', {}).get('value', 'N/A')),
        ('최근 3년 평균', stats.get('최근_3년_평균', {}).get('value', 'N/A')),
        ('연평균 증가율', f"{stats.get('연평균_증가율', {}).get('value', 'N/A')}%"),
        ('변동성 (표준편차)', stats.get('변동성', {}).get('value', 'N/A')),
    ]
    
    for label, value in summary_items:
        para = doc.add_paragraph()
        run = para.add_run(f'• {label}: ')
        run.font.bold = True
        run.font.size = Pt(10)
        para.add_run(str(value)).font.size = Pt(10)
    
    doc.add_page_break()
    
    # 2. 지출목적별 소비자물가지수 동향
    doc.add_heading('2. 지출목적별 소비자물가지수 동향', 1)
    
    if '최고_상승률_지출목적' in stats:
        para = doc.add_paragraph()
        run = para.add_run('최고 상승률 지출목적: ')
        run.font.bold = True
        run.font.size = Pt(10)
        stat = stats['최고_상승률_지출목적']
        para.add_run(f"{stat.get('category', 'N/A')} ({stat.get('value', 'N/A')}%)").font.size = Pt(10)
    
    if '최저_상승률_지출목적' in stats:
        para = doc.add_paragraph()
        run = para.add_run('최저 상승률 지출목적: ')
        run.font.bold = True
        run.font.size = Pt(10)
        stat = stats['최저_상승률_지출목적']
        para.add_run(f"{stat.get('category', 'N/A')} ({stat.get('value', 'N/A')}%)").font.size = Pt(10)
    
    if '상위_지출목적_평균' in stats:
        para = doc.add_paragraph()
        run = para.add_run('상위 지출목적 평균 물가지수: ')
        run.font.bold = True
        run.font.size = Pt(10)
        para.add_run('\n').font.size = Pt(10)
        
        for cat in stats['상위_지출목적_평균'].get('categories', []):
            para2 = doc.add_paragraph()
            run2 = para2.add_run(f"  • {cat['name']}: ")
            run2.font.bold = True
            run2.font.size = Pt(10)
            para2.add_run(f"{cat['value']}").font.size = Pt(10)
    
    doc.add_page_break()
    
    # 3. 종합 분석
    doc.add_heading('3. 종합 분석', 1)
    
    analysis_text = f"""
전체 기간 평균 소비자물가지수는 {stats.get('전체_평균', {}).get('value', 'N/A')}로 나타났으며, 
최근 1년 평균은 {stats.get('최근_1년_평균', {}).get('value', 'N/A')}입니다.

연평균 증가율은 {stats.get('연평균_증가율', {}).get('value', 'N/A')}%로, 
물가가 지속적으로 상승하는 추세를 보이고 있습니다.

최근 6개월 추세는 {stats.get('최근_추세', {}).get('trend', 'N/A')} 추세로, 
이전 6개월 대비 {abs(stats.get('최근_추세', {}).get('value', 0)):.2f}% 변화를 보였습니다.
    """
    
    para = doc.add_paragraph(analysis_text.strip())
    for run in para.runs:
        run.font.size = Pt(10)
    
    # 바이트 스트림으로 반환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/statistics')
def get_statistics():
    try:
        df = load_data()
        stats = calculate_statistics(df)
        return jsonify({'success': True, 'statistics': stats})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_press_release_html(stats):
    """보도자료를 HTML 형식으로 생성"""
    now = datetime.now()
    weekday_kr = ['월', '화', '수', '목', '금', '토', '일'][now.weekday()]
    press_time = f'보도시점 {now.strftime("%Y. %m. %d")}.({weekday_kr}) 08:00'
    release_time = f'배포{now.strftime("%Y. %m. %d")}.({weekday_kr}) 07:30'
    
    # 값 포맷팅 헬퍼 함수
    def format_value(val):
        if val is None:
            return '-'
        try:
            return f"{float(val):.1f}"
        except (ValueError, TypeError):
            return '-'
    
    html = f"""
    <div class="press-release" style="color: #212529;">
        <div class="press-header">
            <h2 style="color: #1a1a1a;">소비자물가조사</h2>
            <h2 style="color: #1a1a1a;">보도자료</h2>
            <p class="press-time" style="color: #495057;">{press_time} {release_time}</p>
            <h1 style="color: #1a1a1a;">{now.strftime("%Y년 %m월")} 소비자물가동향</h1>
            <div class="press-dept" style="color: #495057;">
                <p>담당 부서 경제동향통계심의관 책임자 과  장 박병선(042-481-2530)</p>
                <p>물가동향과 담당자 사무관 이정화(042-481-2531)</p>
            </div>
        </div>
        
        <div class="press-section" style="color: #212529;">
            <h3 style="color: #1a1a1a;">일 러 두 기</h3>
            <ul class="notice-list" style="color: #212529;">
                <li style="color: #212529;"><strong style="color: #667eea;">□</strong> 현재 소비자물가지수의 기준연도는 2020년, 가중치의 기준연도는 2022년입니다.</li>
                <li style="color: #212529;"><strong style="color: #667eea;">○</strong> 따라서 품목별 지수와 가중치를 이용하여 상위 단계 지수 계산한 결과와 공표하는 지수는 일치하지 않음에 유의하여 주시기 바랍니다.</li>
                <li style="color: #212529;">※ 상세내용은 부록 소비자물가지수 계산식 참조</li>
                <li style="color: #212529;"><strong style="color: #667eea;">□</strong> 매월 발표하는 소비자물가지수는 가격변동을 측정하는 것으로 가격의 절대수준을 나타내지 않습니다.</li>
                <li style="color: #212529;"><strong style="color: #667eea;">○</strong> 따라서 지역별로 기준시점(2020년=100)의 가격수준이 다르기 때문에 지역별 소비자물가지수를 이용하여 지역간 상대적인 물가수준 차이를 비교하는 것은 부적절합니다.</li>
                <li style="color: #212529;"><strong style="color: #667eea;">□</strong> 일반적으로 소비자물가변동 추이 및 국가 간 비교는 1년 전 대비 물가 변동인 전년동월비를 주로 이용하지만, 단기간의 변동인 전월비도 참고하시기 바랍니다.</li>
                <li style="color: #212529;"><strong style="color: #667eea;">□</strong> 소비자물가지수는 2019년 이전은 소수점 이하 3자리, 2020년 이후는 소수점 이하 2자리로 작성되고 있습니다.</li>
                <li style="color: #212529;"><strong style="color: #667eea;">○</strong> 통계표에 사용된 "-" 부호의 뜻은 "해당 숫자 없음"을 의미합니다.</li>
                <li style="color: #212529;"><strong style="color: #667eea;">□</strong> 본문에 수록된 자료는 국가데이터처 홈페이지(http://kostat.go.kr) 및 국가통계포털(http://kosis.kr)을 통해 이용할 수 있습니다.</li>
            </ul>
        </div>
        
        <div class="press-section" style="color: #212529;">
            <h3 style="color: #1a1a1a;">1. 주요 통계량 분석</h3>
            
            <h4 style="color: #2d2d2d;">핵심 통계 지표</h4>
            <ul class="stats-summary" style="color: #212529;">
    """
    
    # 1. 최근 3개월 평균 증가율
    recent_3m = stats.get('최근_3개월_평균_증가율', {})
    if recent_3m.get('value') is not None:
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>최근 3개월 평균 증가율:</strong> {format_value(recent_3m.get('value'))}% (전년동월비)</li>"
    
    # 2. 최고 상승률 달
    max_month = stats.get('최고_상승률_달', {})
    if max_month.get('value') is not None:
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>최고 상승률 달:</strong> {max_month.get('date', 'N/A')} ({format_value(max_month.get('value'))}%)</li>"
    
    # 3. 최저 상승률 달
    min_month = stats.get('최저_상승률_달', {})
    if min_month.get('value') is not None:
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>최저 상승률 달:</strong> {min_month.get('date', 'N/A')} ({format_value(min_month.get('value'))}%)</li>"
    
    # 4. 물가 상승 추세
    trend = stats.get('물가_상승_추세', {})
    if trend.get('value') is not None:
        trend_val = trend.get('value')
        trend_desc = '상승' if trend_val > 0 else '하락' if trend_val < 0 else '유지'
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>물가 상승 추세:</strong> {format_value(trend_val)} (월당, {trend_desc} 추세)</li>"
    
    # 5. 변동성 지수
    volatility = stats.get('변동성_지수', {})
    if volatility.get('value') is not None:
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>변동성 지수:</strong> {format_value(volatility.get('value'))} (표준편차)</li>"
    
    # 6. 최고 변동성 지출목적
    max_vol = stats.get('최고_변동성_지출목적', {})
    if max_vol.get('category'):
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>최고 변동성 지출목적:</strong> {max_vol.get('category', 'N/A')} (변동성: {format_value(max_vol.get('value'))})</li>"
    
    # 7. 최저 변동성 지출목적
    min_vol = stats.get('최저_변동성_지출목적', {})
    if min_vol.get('category'):
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>최저 변동성 지출목적:</strong> {min_vol.get('category', 'N/A')} (변동성: {format_value(min_vol.get('value'))})</li>"
    
    # 8. 물가 안정성 점수
    stability = stats.get('물가_안정성_점수', {})
    if stability.get('value') is not None:
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>물가 안정성 점수:</strong> {format_value(stability.get('value'))}점 (0-100, 높을수록 안정적)</li>"
    
    # 9. 최근 6개월 변화
    recent_6m = stats.get('최근_6개월_변화', {})
    if recent_6m.get('value') is not None:
        trend_desc = recent_6m.get('trend', '유지')
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>최근 6개월 변화:</strong> {format_value(recent_6m.get('value'))}% ({trend_desc} 추세)</li>"
    
    # 10. 계절성 패턴
    seasonal = stats.get('계절성_패턴', {})
    if seasonal.get('highest_month'):
        html += f"<li style='color: #212529;'><strong style='color: #1a1a1a;'>계절성 패턴:</strong> 최고 {seasonal.get('highest_month', 'N/A')} ({format_value(seasonal.get('highest_value'))}), 최저 {seasonal.get('lowest_month', 'N/A')} ({format_value(seasonal.get('lowest_value'))})</li>"
    
    html += """
            </ul>
        </div>
        
        <div class="press-section" style="color: #212529;">
            <h3 style="color: #1a1a1a;">2. 지출목적별 분석</h3>
    """
    
    if '최고_변동성_지출목적' in stats:
        max_vol = stats['최고_변동성_지출목적']
        html += f"<h4 style='color: #2d2d2d;'>가장 변동성이 큰 지출목적</h4>"
        html += f"<p style='color: #212529;'><strong style='color: #1a1a1a;'>{max_vol.get('category', 'N/A')}:</strong> 변동성 {format_value(max_vol.get('value'))}</p>"
    
    if '최저_변동성_지출목적' in stats:
        min_vol = stats['최저_변동성_지출목적']
        html += f"<h4 style='color: #2d2d2d;'>가장 안정적인 지출목적</h4>"
        html += f"<p style='color: #212529;'><strong style='color: #1a1a1a;'>{min_vol.get('category', 'N/A')}:</strong> 변동성 {format_value(min_vol.get('value'))}</p>"
    
    html += """
        </div>
        
        <div class="press-section" style="color: #212529;">
            <h3 style="color: #1a1a1a;">3. 종합 분석</h3>
            <div class="analysis-content" style="color: #212529;">
    """
    
    # 종합 분석 텍스트 생성
    analysis_parts = []
    
    if recent_3m.get('value') is not None:
        analysis_parts.append(f"최근 3개월 평균 증가율은 {format_value(recent_3m.get('value'))}%로 나타났습니다.")
    
    if trend.get('value') is not None:
        trend_desc = '상승' if trend.get('value') > 0 else '하락' if trend.get('value') < 0 else '유지'
        analysis_parts.append(f"물가 상승 추세는 {trend_desc} 추세를 보이고 있습니다.")
    
    if stability.get('value') is not None:
        stability_val = float(stability.get('value'))
        if stability_val >= 70:
            stability_desc = "매우 안정적"
        elif stability_val >= 50:
            stability_desc = "안정적"
        else:
            stability_desc = "불안정"
        analysis_parts.append(f"물가 안정성 점수는 {format_value(stability.get('value'))}점으로 {stability_desc}입니다.")
    
    if seasonal.get('highest_month'):
        analysis_parts.append(f"계절성 패턴을 보면 {seasonal.get('highest_month')}에 가장 높고 {seasonal.get('lowest_month')}에 가장 낮은 물가지수를 보입니다.")
    
    if recent_6m.get('value') is not None:
        trend_desc = recent_6m.get('trend', '유지')
        analysis_parts.append(f"최근 6개월은 이전 6개월 대비 {format_value(recent_6m.get('value'))}% {trend_desc} 추세를 보였습니다.")
    
    analysis_text = ""
    for part in analysis_parts:
        analysis_text += f"<p style='color: #212529;'>{part}</p>\n                "
    
    if not analysis_text:
        analysis_text = "<p style='color: #212529;'>통계량 데이터를 분석한 결과를 종합적으로 제시합니다.</p>"
    
    html += analysis_text
    html += """
            </div>
        </div>
    </div>
    """
    
    return html

@app.route('/api/press-release')
def get_press_release():
    """보도자료를 HTML 형식으로 반환"""
    try:
        df = load_data()
        stats = calculate_statistics(df)
        html_content = generate_press_release_html(stats)
        return jsonify({'success': True, 'html': html_content})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=8889)

