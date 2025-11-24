from flask import Flask, render_template, jsonify, send_file, request
import pandas as pd
import numpy as np
from datetime import datetime
import io
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()

def load_data():
    """Excel 파일 로드"""
    file_path = '지출목적별_소비자물가지수_품목포함__2020100__20251106131304.xlsx'
    df = pd.read_excel(file_path, sheet_name='데이터')
    return df

def calculate_statistics(df):
    """주요 통계량 10개 계산"""
    # 데이터 전처리
    data_cols = [col for col in df.columns if col not in ['시도별', '지출목적별']]
    numeric_data = df[data_cols].apply(pd.to_numeric, errors='coerce')
    
    # 전체 데이터 (총지수만)
    total_index = df[df['지출목적별'] == '0 총지수'].iloc[0]
    total_values = pd.to_numeric(total_index[data_cols], errors='coerce').dropna()
    
    stats = {}
    
    # 1. 전체 평균 물가지수 (2020=100 기준)
    stats['전체_평균'] = {
        'value': round(total_values.mean(), 2),
        'description': '전체 기간 평균 소비자물가지수'
    }
    
    # 2. 최고/최저 물가지수
    stats['최고_물가지수'] = {
        'value': round(total_values.max(), 2),
        'date': total_values.idxmax(),
        'description': '전체 기간 중 최고 물가지수'
    }
    stats['최저_물가지수'] = {
        'value': round(total_values.min(), 2),
        'date': total_values.idxmin(),
        'description': '전체 기간 중 최저 물가지수'
    }
    
    # 3. 최근 1년 평균 (마지막 12개월)
    recent_cols = data_cols[-12:] if len(data_cols) >= 12 else data_cols
    recent_values = pd.to_numeric(total_index[recent_cols], errors='coerce').dropna()
    stats['최근_1년_평균'] = {
        'value': round(recent_values.mean(), 2),
        'description': '최근 1년(2024.11~2025.10) 평균 물가지수'
    }
    
    # 4. 최근 3년 평균 (2022.11 ~ 2025.10)
    three_year_cols = [col for col in data_cols if col >= '2022.11']
    three_year_values = pd.to_numeric(total_index[three_year_cols], errors='coerce').dropna()
    stats['최근_3년_평균'] = {
        'value': round(three_year_values.mean(), 2),
        'description': '최근 3년 평균 물가지수'
    }
    
    # 5. 연도별 평균 증가율
    years = sorted(set([col.split('.')[0] for col in data_cols if '.' in col]))
    year_avgs = {}
    for year in years:
        year_cols = [col for col in data_cols if col.startswith(year)]
        year_values = pd.to_numeric(total_index[year_cols], errors='coerce').dropna()
        if len(year_values) > 0:
            year_avgs[year] = year_values.mean()
    
    if len(year_avgs) >= 2:
        years_sorted = sorted(year_avgs.keys())
        growth_rates = []
        for i in range(1, len(years_sorted)):
            prev_avg = year_avgs[years_sorted[i-1]]
            curr_avg = year_avgs[years_sorted[i]]
            if prev_avg > 0:
                growth_rate = ((curr_avg - prev_avg) / prev_avg) * 100
                growth_rates.append(growth_rate)
        
        avg_growth_rate = np.mean(growth_rates) if growth_rates else 0
        stats['연평균_증가율'] = {
            'value': round(avg_growth_rate, 2),
            'unit': '%',
            'description': '연평균 물가 상승률'
        }
    
    # 6. 변동성 (표준편차)
    stats['변동성'] = {
        'value': round(total_values.std(), 2),
        'description': '전체 기간 물가지수 표준편차 (변동성)'
    }
    
    # 7. 최고 상승률 지출목적
    category_growth = {}
    for idx, row in df.iterrows():
        if pd.isna(row['지출목적별']) or row['지출목적별'] == '0 총지수':
            continue
        
        cat_values = pd.to_numeric(row[data_cols], errors='coerce').dropna()
        if len(cat_values) >= 2:
            first_val = cat_values.iloc[0]
            last_val = cat_values.iloc[-1]
            if first_val > 0:
                growth = ((last_val - first_val) / first_val) * 100
                category_growth[row['지출목적별']] = growth
    
    if category_growth:
        max_category = max(category_growth.items(), key=lambda x: x[1])
        stats['최고_상승률_지출목적'] = {
            'category': max_category[0],
            'value': round(max_category[1], 2),
            'unit': '%',
            'description': '전체 기간 중 가장 높은 상승률을 보인 지출목적'
        }
    
    # 8. 최저 상승률 지출목적
    if category_growth:
        min_category = min(category_growth.items(), key=lambda x: x[1])
        stats['최저_상승률_지출목적'] = {
            'category': min_category[0],
            'value': round(min_category[1], 2),
            'unit': '%',
            'description': '전체 기간 중 가장 낮은 상승률을 보인 지출목적'
        }
    
    # 9. 지출목적별 평균 물가지수 (상위 3개)
    category_avgs = {}
    for idx, row in df.iterrows():
        if pd.isna(row['지출목적별']) or row['지출목적별'] == '0 총지수':
            continue
        
        cat_values = pd.to_numeric(row[data_cols], errors='coerce').dropna()
        if len(cat_values) > 0:
            category_avgs[row['지출목적별']] = cat_values.mean()
    
    if category_avgs:
        top_categories = sorted(category_avgs.items(), key=lambda x: x[1], reverse=True)[:3]
        stats['상위_지출목적_평균'] = {
            'categories': [{'name': cat[0], 'value': round(cat[1], 2)} for cat in top_categories],
            'description': '평균 물가지수가 가장 높은 지출목적 상위 3개'
        }
    
    # 10. 추세 분석 (최근 6개월 vs 이전 6개월)
    if len(data_cols) >= 12:
        recent_6m = [col for col in data_cols[-6:]]
        prev_6m = [col for col in data_cols[-12:-6]]
        
        recent_6m_values = pd.to_numeric(total_index[recent_6m], errors='coerce').dropna()
        prev_6m_values = pd.to_numeric(total_index[prev_6m], errors='coerce').dropna()
        
        if len(recent_6m_values) > 0 and len(prev_6m_values) > 0:
            recent_avg = recent_6m_values.mean()
            prev_avg = prev_6m_values.mean()
            trend_change = ((recent_avg - prev_avg) / prev_avg) * 100 if prev_avg > 0 else 0
            
            stats['최근_추세'] = {
                'value': round(trend_change, 2),
                'unit': '%',
                'description': '최근 6개월 평균 대비 이전 6개월 대비 변화율',
                'trend': '상승' if trend_change > 0 else '하락' if trend_change < 0 else '유지'
            }
    
    return stats

def generate_press_release(stats):
    """보도자료 생성"""
    doc = Document()
    
    # 제목
    title = doc.add_heading('지출목적별 소비자물가지수 주요 통계 분석', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 날짜
    date_para = doc.add_paragraph(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 빈 줄
    
    # 개요
    doc.add_heading('1. 분석 개요', 1)
    doc.add_paragraph(
        '본 보고서는 지출목적별 소비자물가지수 데이터를 분석하여 주요 통계량 10개를 도출하고, '
        '물가 동향을 종합적으로 분석한 결과를 제시합니다.'
    )
    
    doc.add_paragraph()  # 빈 줄
    
    # 주요 통계량
    doc.add_heading('2. 주요 통계량', 1)
    
    stat_order = [
        '전체_평균', '최고_물가지수', '최저_물가지수', '최근_1년_평균', 
        '최근_3년_평균', '연평균_증가율', '변동성', '최고_상승률_지출목적',
        '최저_상승률_지출목적', '상위_지출목적_평균', '최근_추세'
    ]
    
    for i, stat_key in enumerate(stat_order, 1):
        if stat_key in stats:
            stat = stats[stat_key]
            doc.add_heading(f'2.{i} {stat["description"]}', 2)
            
            para = doc.add_paragraph()
            if 'value' in stat:
                value_text = f'{stat["value"]}'
                if 'unit' in stat:
                    value_text += f' {stat["unit"]}'
                para.add_run(f'값: ').bold = True
                para.add_run(value_text)
            
            if 'date' in stat:
                para = doc.add_paragraph()
                para.add_run(f'시점: ').bold = True
                para.add_run(stat['date'])
            
            if 'category' in stat:
                para = doc.add_paragraph()
                para.add_run(f'지출목적: ').bold = True
                para.add_run(stat['category'])
            
            if 'categories' in stat:
                for cat in stat['categories']:
                    para = doc.add_paragraph()
                    para.add_run(f'• {cat["name"]}: ').bold = True
                    para.add_run(f'{cat["value"]}')
            
            if 'trend' in stat:
                para = doc.add_paragraph()
                para.add_run(f'추세: ').bold = True
                para.add_run(stat['trend'])
            
            doc.add_paragraph()  # 빈 줄
    
    # 종합 분석
    doc.add_heading('3. 종합 분석', 1)
    
    analysis_text = f"""
    전체 기간 평균 소비자물가지수는 {stats.get('전체_평균', {}).get('value', 'N/A')}로 나타났으며, 
    최근 1년 평균은 {stats.get('최근_1년_평균', {}).get('value', 'N/A')}입니다.
    
    연평균 증가율은 {stats.get('연평균_증가율', {}).get('value', 'N/A')}%로, 
    물가가 지속적으로 상승하는 추세를 보이고 있습니다.
    
    최근 6개월 추세는 {stats.get('최근_추세', {}).get('trend', 'N/A')} 추세로, 
    이전 6개월 대비 {abs(stats.get('최근_추세', {}).get('value', 0))}% 변화를 보였습니다.
    """
    
    doc.add_paragraph(analysis_text.strip())
    
    # 바이트 스트림으로 반환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/statistics')
def get_statistics():
    try:
        df = load_data()
        stats = calculate_statistics(df)
        return jsonify({'success': True, 'statistics': stats})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/press-release')
def download_press_release():
    try:
        df = load_data()
        stats = calculate_statistics(df)
        buffer = generate_press_release(stats)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'소비자물가지수_보도자료_{datetime.now().strftime("%Y%m%d")}.docx'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def extract_text_from_pdf(pdf_file):
    """PDF 파일에서 텍스트 추출"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = []
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text.strip():
                text_content.append(text)
        
        return '\n\n'.join(text_content)
    except Exception as e:
        raise Exception(f"PDF 텍스트 추출 실패: {str(e)}")

def convert_pdf_to_docx(pdf_file):
    """PDF를 DOCX로 변환 (pypandoc 사용)"""
    if not PYPANDOC_AVAILABLE:
        raise Exception("pypandoc이 설치되지 않았습니다. 'pip install pypandoc'으로 설치해주세요.")
    
    if not PYPDF2_AVAILABLE:
        raise Exception("PyPDF2가 설치되지 않았습니다. 'pip install PyPDF2'로 설치해주세요.")
    
    try:
        # PDF에서 텍스트 추출
        text_content = extract_text_from_pdf(pdf_file)
        
        if not text_content.strip():
            raise Exception("PDF에서 텍스트를 추출할 수 없습니다. 이미지 기반 PDF일 수 있습니다.")
        
        # 임시 마크다운 파일 생성
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as md_file:
            md_file.write(text_content)
            md_file_path = md_file.name
        
        docx_file_path = None
        try:
            # 임시 DOCX 파일 생성
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
                docx_file_path = docx_file.name
            
            # pypandoc을 사용하여 마크다운을 DOCX로 변환
            try:
                pypandoc.convert_file(
                    md_file_path,
                    'docx',
                    outputfile=docx_file_path,
                    format='markdown'
                )
            except RuntimeError as e:
                error_msg = str(e)
                if 'pandoc' in error_msg.lower() or 'not found' in error_msg.lower():
                    raise Exception("Pandoc이 시스템에 설치되지 않았습니다. Pandoc을 설치해주세요. (macOS: brew install pandoc, Ubuntu: sudo apt-get install pandoc)")
                else:
                    raise Exception(f"Pandoc 변환 오류: {error_msg}")
            
            # 변환된 DOCX 파일 읽기
            with open(docx_file_path, 'rb') as f:
                docx_content = f.read()
            
            # 임시 파일 삭제
            os.unlink(md_file_path)
            os.unlink(docx_file_path)
            
            return docx_content
            
        except Exception as e:
            # 에러 발생 시 임시 파일 정리
            if os.path.exists(md_file_path):
                os.unlink(md_file_path)
            if docx_file_path and os.path.exists(docx_file_path):
                os.unlink(docx_file_path)
            raise
            
    except Exception as e:
        raise Exception(f"PDF to DOCX 변환 실패: {str(e)}")

@app.route('/api/pdf-to-docx', methods=['POST'])
def pdf_to_docx():
    """PDF 파일을 DOCX로 변환"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '파일이 업로드되지 않았습니다.'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'success': False, 'error': 'PDF 파일만 업로드 가능합니다.'}), 400
        
        # PDF를 DOCX로 변환
        docx_content = convert_pdf_to_docx(file)
        
        # 원본 파일명에서 확장자 변경
        original_filename = file.filename.rsplit('.', 1)[0]
        output_filename = f'{original_filename}.docx'
        
        # 바이트 스트림으로 반환
        buffer = io.BytesIO(docx_content)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        )
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=8889)

